# ============================
#  Final Version – Minutes of Meeting Exporter
# ============================

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    Image as RLImage,
)
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from transformers import BartTokenizer, BartForConditionalGeneration
from datetime import datetime
from io import BytesIO
import re
import os


class MeetingExporter:
    def __init__(self, header_image_path="college_header.jpg"):
        """
        header_image_path: The banner/header image for the top of PDF & DOCX.
        """
        self.header_image_path = header_image_path

        # Load BART model (only once)
        self.tokenizer = BartTokenizer.from_pretrained("facebook/bart-large-cnn")
        self.model = BartForConditionalGeneration.from_pretrained("facebook/bart-large-cnn")

    # --------------------------
    # BASIC CLEANERS
    # --------------------------

    @staticmethod
    def _sanitize(text: str) -> str:
        """Remove strange characters, normalize spacing."""
        if not text:
            return ""
        text = re.sub(r"[\u2580-\u259F\u2500-\u257F\u25A0-\u25FF]+", "-", text)
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    @staticmethod
    def _clean_action_text(text: str) -> str:
        """Light cleanup for task text before NLP rewriting."""
        if not text:
            return ""
        text = re.sub(r"\b(ma|ma\.|am|am\.)\b", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s+", " ", text)
        return text.strip(" .,-")

    # --------------------------
    # NLP REWRITING
    # --------------------------

    def _rewrite_action_sentence(self, task, responsible="", deadline=""):
        """
        Converts task+responsible+deadline into a CLEAN PROFESSIONAL SENTENCE.
        No status.
        Example: "John will prepare the budget report by Monday."
        """

        prompt = (
            "Rewrite the following action item into a single clear, formal, "
            "professional sentence appropriate for college meeting minutes.\n\n"
            f"Task: {task}\n"
        )

        if responsible:
            prompt += f"Responsible: {responsible}\n"
        if deadline:
            prompt += f"Deadline: {deadline}\n"

        inputs = self.tokenizer(prompt, return_tensors="pt", truncation=True, max_length=512)

        summary_ids = self.model.generate(
            inputs["input_ids"],
            num_beams=4,
            max_length=40,
            min_length=10,
            no_repeat_ngram_size=3,
            early_stopping=True
        )

        sentence = self.tokenizer.decode(summary_ids[0], skip_special_tokens=True)
        return sentence.strip()

    def _rewrite_summary_formal(self, summary_text: str) -> str:
        cleaned = self._sanitize(summary_text)
        if not cleaned:
            return ""

        # Split into sentences, clean and rejoin nicely
        sentences = [s.strip().capitalize() for s in re.split(r'(?<=[.!?])\s+', cleaned) if s.strip()]
        joined = " ".join(sentences)

        formal = f"The meeting was convened to discuss the following points: {joined}"
        return formal

    # --------------------------
    # HEADER IMAGE HELPERS
    # --------------------------

    def _add_docx_header(self, doc: Document):
        """Adds header image to DOCX if available."""
        if os.path.exists(self.header_image_path):
            try:
                img = doc.add_picture(self.header_image_path, width=Inches(6.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

    def _add_pdf_header(self, story):
        """Adds header image to PDF if available."""
        if os.path.exists(self.header_image_path):
            try:
                img = RLImage(self.header_image_path, width=6.5 * inch)
                story.append(img)
                story.append(Spacer(1, 12))
            except Exception:
                pass

# ============================
#  DOCX EXPORT SECTION
# ============================

    # ---------------------------------------------------------
    #                    DOCX EXPORT
    # ---------------------------------------------------------

    def export_to_docx(self, meeting_data):
        """
        Generates a DOCX Minutes of Meeting document with:
        - Header image
        - Formal summary
        - Agenda table
        - Rewritten natural action items
        """
        doc = Document()

        # Add header image
        self._add_docx_header(doc)

        # Title
        title = doc.add_heading("Minutes of Meeting", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # ----------------------------
        # METADATA
        # ----------------------------
        metadata = meeting_data.get("metadata", {})

        def add_meta(label, key):
            if metadata.get(key):
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(metadata[key]))

        add_meta("Title", "title")
        add_meta("Date", "date")
        add_meta("Time", "time")
        add_meta("Venue", "venue")
        add_meta("Organizer", "organizer")
        add_meta("Recorder", "recorder")

        doc.add_paragraph("----")

        # ----------------------------
        # ATTENDEES
        # ----------------------------
        attendees = meeting_data.get("attendees", [])
        if attendees:
            doc.add_heading("Attendees", level=2)
            for person in attendees:
                name = self._sanitize(person.get("name", ""))
                role = self._sanitize(person.get("role", ""))
                if role:
                    doc.add_paragraph(f"{name} – {role}", style="List Bullet")
                else:
                    doc.add_paragraph(name, style="List Bullet")
            doc.add_paragraph("----")

        # ----------------------------
        # AGENDA TABLE
        # ----------------------------
        agenda = meeting_data.get("agenda", [])
        if agenda:
            doc.add_heading("Agenda", level=2)

            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"

            hdr = table.rows[0].cells
            hdr[0].text = "Agenda No."
            hdr[1].text = "Discussion & Action to be taken"
            hdr[2].text = "Responsibility"

            # bold header
            for cell in hdr:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True

            # rows
            for item in agenda:
                row = table.add_row().cells
                row[0].text = str(item.get("no", ""))
                row[1].text = self._sanitize(item.get("discussion", ""))
                row[2].text = self._sanitize(item.get("responsibility", ""))

            doc.add_paragraph("----")

        # ----------------------------
        # DISCUSSION SUMMARY  (FORMAL REWRITE)
        # ----------------------------
        summary = meeting_data.get("summary", "")
        if summary:
            doc.add_heading("Discussion Summary", level=2)

            rewritten = self._rewrite_summary_formal(summary)
            doc.add_paragraph(rewritten)

            doc.add_paragraph("----")

        # ----------------------------
        # DECISIONS
        # ----------------------------
        decisions = meeting_data.get("decisions", [])
        if decisions:
            doc.add_heading("Decisions", level=2)
            for d in decisions:
                doc.add_paragraph(self._sanitize(d), style="List Bullet")
            doc.add_paragraph("----")

        # ----------------------------
        # ACTION ITEMS (Formal rewritten sentences, no status)
        # ----------------------------
        action_items = meeting_data.get("action_items", [])
        if action_items:
            doc.add_heading("Action Items", level=2)

            for action in action_items:
                task = self._clean_action_text(self._sanitize(action.get("task", "")))
                responsible = self._sanitize(action.get("responsible", ""))
                deadline = self._sanitize(action.get("deadline", ""))

                rewritten = self._rewrite_action_sentence(task, responsible, deadline)
                doc.add_paragraph(f"• {rewritten}")

            doc.add_paragraph("----")

        # ----------------------------
        # NEXT MEETING
        # ----------------------------
        next_meeting = meeting_data.get("next_meeting", {})
        if any(next_meeting.values()):
            doc.add_heading("Next Meeting", level=2)

            def nm(label, key):
                if next_meeting.get(key):
                    p = doc.add_paragraph()
                    p.add_run(f"{label}: ").bold = True
                    p.add_run(str(next_meeting[key]))

            nm("Date", "date")
            nm("Time", "time")
            nm("Venue", "venue")
            nm("Agenda", "agenda")

            doc.add_paragraph("─" * 50)

        # ----------------------------
        # CLOSING NOTE
        # ----------------------------
        doc.add_heading("Closing Note", level=2)
        closing = (
            "This document was generated by the AIMS – AI Meeting Summarizer "
            f"on {datetime.now().strftime('%d/%m/%Y at %H:%M')}."
        )
        doc.add_paragraph(closing)

        # save
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

# ============================
#  export_utils.py (PART 3)
#  PDF EXPORT SECTION
# ============================

    # ---------------------------------------------------------
    #                    PDF EXPORT
    # ---------------------------------------------------------

    def export_to_pdf(self, meeting_data):
        """
        Generates a PDF Minutes of Meeting document with:
        - Header image
        - Agenda table
        - Formal rewritten summary
        - Rewritten natural action items
        """
        buffer = BytesIO()

        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=20
        )

        # ----------------------------
        # STYLES
        # ----------------------------
        styles = getSampleStyleSheet()

        styles.add(ParagraphStyle(
            name="TitleStyle",
            fontSize=20,
            leading=24,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#1a237e"),
            spaceAfter=30
        ))

        styles.add(ParagraphStyle(
            name="HeadingStyle",
            parent=styles["Heading2"],
            fontSize=14,
            leading=16,
            textColor=colors.HexColor("#283593"),
            spaceAfter=12,
            spaceBefore=12
        ))

        styles.add(ParagraphStyle(
            name="BodyStyle",
            parent=styles["BodyText"],
            fontSize=11,
            leading=14,
            spaceAfter=6
        ))

        styles.add(ParagraphStyle(
            name="TableBody",
            parent=styles["BodyText"],
            fontSize=10,
            leading=12
        ))

        story = []

        # ----------------------------
        # HEADER IMAGE
        # ----------------------------
        self._add_pdf_header(story)

        # ----------------------------
        # TITLE
        # ----------------------------
        story.append(Paragraph("Minutes of Meeting", styles["TitleStyle"]))
        story.append(Spacer(1, 12))

        # ----------------------------
        # METADATA
        # ----------------------------
        metadata = meeting_data.get("metadata", {})

        def add_meta(label, key):
            if metadata.get(key):
                story.append(Paragraph(f"<b>{label}:</b> {metadata[key]}", styles["BodyStyle"]))

        add_meta("Title", "title")
        add_meta("Date", "date")
        add_meta("Time", "time")
        add_meta("Venue", "venue")
        add_meta("Organizer", "organizer")
        add_meta("Recorder", "recorder")

        story.append(Spacer(1, 12))
        story.append(Paragraph("----", styles["BodyStyle"]))
        story.append(Spacer(1, 12))

        # ----------------------------
        # ATTENDEES
        # ----------------------------
        attendees = meeting_data.get("attendees", [])
        if attendees:
            story.append(Paragraph("Attendees", styles["HeadingStyle"]))
            for person in attendees:
                name = self._sanitize(person.get("name", ""))
                role = self._sanitize(person.get("role", ""))
                if role:
                    story.append(Paragraph(f"• {name} – {role}", styles["BodyStyle"]))
                else:
                    story.append(Paragraph(f"• {name}", styles["BodyStyle"]))

            story.append(Spacer(1, 12))
            story.append(Paragraph("----", styles["BodyStyle"]))
            story.append(Spacer(1, 12))

        # ----------------------------
        # AGENDA TABLE
        # ----------------------------
        agenda = meeting_data.get("agenda", [])
        if agenda:
            story.append(Paragraph("Agenda", styles["HeadingStyle"]))

            table_data = [
                ["Agenda No.", "Discussion & Action to be taken", "Responsibility"]
            ]

            for item in agenda:
                table_data.append([
                    str(item.get("no", "")),
                    self._sanitize(item.get("discussion", "")),
                    self._sanitize(item.get("responsibility", ""))
                ])

            table = Table(
                table_data,
                colWidths=[60, 330, 110]
            )

            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EAF6")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 10),
                ("GRID", (0, 0), (-1, -1), 0.7, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 1), (-1, -1), 10),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]))

            story.append(table)
            story.append(Spacer(1, 18))
            story.append(Paragraph("----", styles["BodyStyle"]))
            story.append(Spacer(1, 12))

        # ----------------------------
        # DISCUSSION SUMMARY (FORMAL REWRITE)
        # ----------------------------
        summary = meeting_data.get("summary", "")
        if summary:
            story.append(Paragraph("Discussion Summary", styles["HeadingStyle"]))

            rewritten = self._rewrite_summary_formal(summary)
            story.append(Paragraph(rewritten, styles["BodyStyle"]))

            story.append(Spacer(1, 12))
            story.append(Paragraph("----", styles["BodyStyle"]))
            story.append(Spacer(1, 12))

        # ----------------------------
        # DECISIONS
        # ----------------------------
        decisions = meeting_data.get("decisions", [])
        if decisions:
            story.append(Paragraph("Decisions", styles["HeadingStyle"]))
            for dec in decisions:
                story.append(Paragraph(f"• {self._sanitize(dec)}", styles["BodyStyle"]))

            story.append(Spacer(1, 12))
            story.append(Paragraph("----", styles["BodyStyle"]))
            story.append(Spacer(1, 12))

        # ----------------------------
        # ACTION ITEMS (Rewritten sentences)
        # ----------------------------
        action_items = meeting_data.get("action_items", [])
        if action_items:
            story.append(Paragraph("Action Items", styles["HeadingStyle"]))

            for action in action_items:
                task = self._clean_action_text(self._sanitize(action.get("task", "")))
                responsible = self._sanitize(action.get("responsible", ""))
                deadline = self._sanitize(action.get("deadline", ""))

                rewritten = self._rewrite_action_sentence(task, responsible, deadline)
                story.append(Paragraph(f"• {rewritten}", styles["BodyStyle"]))

            story.append(Spacer(1, 12))
            story.append(Paragraph("----", styles["BodyStyle"]))
            story.append(Spacer(1, 12))

        # ----------------------------
        # NEXT MEETING
        # ----------------------------
        next_meeting = meeting_data.get("next_meeting", {})
        if any(next_meeting.values()):
            story.append(Paragraph("Next Meeting", styles["HeadingStyle"]))

            def nm(label, key):
                if next_meeting.get(key):
                    story.append(
                        Paragraph(f"<b>{label}:</b> {next_meeting[key]}", styles["BodyStyle"])
                    )

            nm("Date", "date")
            nm("Time", "time")
            nm("Venue", "venue")
            nm("Agenda", "agenda")

            story.append(Spacer(1, 12))
            story.append(Paragraph("─" * 90, styles["BodyStyle"]))
            story.append(Spacer(1, 12))

        # ----------------------------
        # CLOSING NOTE
        # ----------------------------
        story.append(Paragraph("Closing Note", styles["HeadingStyle"]))

        closing = (
            "This document was generated by the AIMS – AI Meeting Summarizer "
            f"on {datetime.now().strftime('%d/%m/%Y at %H:%M')}."
        )
        story.append(Paragraph(closing, styles["BodyStyle"]))

        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
